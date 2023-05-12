from django.shortcuts import render
from django.views import View
# формы
from .forms import contactForm, ChangeNumberUser, ChangeEmailUser, ChangeAvatarUser, ChangePasswordUser
# модели
from .models import inst, complex, CustomUser, test, question, answers, res, downloadInstructionsForTests, typeuser, Groups
# варианты ответов
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
# обязательная авторизация
from django.contrib.auth.decorators import login_required
# получение времени
from django.utils import timezone 
import timedelta # чтобы переводить даты
# чат бот
import datetime
# импорт ворд
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # выравнивание
from docx.enum.section import WD_ORIENTATION  # jhb
from docx.shared import Pt  # размер шрифта
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.shared import Cm

# чат-бот


# def chatbot_responseView(request):

#     page_name = 'bot.html'  # страница

#     if request.method == 'POST':
#         data = {}
#         query = request.POST
#         user_text = query['message']
#         model = build_model('kbqa_cq_ru', download=True)
#         bot_response = model([user_text])
#         data['status'] = 'ok'
#         return JsonResponse({'data': data, 'response': bot_response[0]})

#     return render(request, page_name)


#главная
class MainView(View):

    def showPage(request):

        page_name = 'main.html'

        users = CustomUser.get_count_users()  # количество пользователей
        # количество пройденных инструктажей
        passed_brief = res.objects.filter(mark='Сдан').count() or 0
        count_brief = test.objects.all().count() or 0  # количество инструктажей

        getInstr = inst.get_all_names()
        getComplex = complex.get_all_names()
        users = CustomUser.get_count_users()
        contact = contactForm()

        values = {
            'contact': contact,
            'instr': getInstr,
            'complex': getComplex,
            'countUser': users,
            'pb': passed_brief,
            'cb': count_brief,
        }

        return render(request, page_name, values)


    def mainSendMessage(request):
        if request.method == 'POST':
            data = {}
            print(request.POST)
            form = contactForm(request.POST)
            if form.is_valid():
                form.save()
                data['message'] = 'Вы успешно отправили!'
                data['status'] = 'ok'
                return JsonResponse(data)
            else:
                data['message'] = 'Введите правильные данные!'
                data['status'] = 'error'
                return JsonResponse(data)


    # подробности
    def mainAboutPage(request):
        page_name = 'about.html'
        obj = CustomUser.objects.get(username='admin')
        return render(request, page_name, {'user':obj})

    # контакты
    def mainContactPage(request):
        page_name = 'contact.html'
        return render(request, page_name)


class JournalView(View):

    def showPage(request):
        pass

    pass


@login_required(login_url='account/login/')
def journalView(request):

    page_name = 'forOss/journal.html'

    if request.user.is_superuser or str(request.user.role).casefold() == 'специалист по охране труда':

        obj_res = None  # Для результатов
        obj_briefs = inst.objects.all()  # получение инструктажей
        obj_type_users = typeuser.objects.all()  # получение тип пользователей
        obj_groups = Groups.objects.all()  # получение групп
        obj_quizes = test.objects.all()  # получение тестов
        obj_results_user = res.objects.all() # получение зачетов

        if 'import-doc' in request.POST:  # условие для скачивания
            print(request.POST)
            obj_result = None

            if request.POST['date_start'] != '':  # Если есть даты
                print(request.POST)
                try:
                    obj_result = res.objects.filter(
                        user__type_user=request.POST['type_user'],
                        instruction=request.POST['type_brief'],
                        user__groupStud_id=request.POST['group'],
                        quiz__id=request.POST['quiz'],
                        date_instruction__range=(request.POST['date_start'],
                                                 timezone.now()),
                        mark=request.POST['mark'],
                    )

                except res.DoesNotExist:
                    obj_result = []

            # styles for docs
            alignment_dict = {
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            }

            orient_dict = {
                'portrait': WD_ORIENT.PORTRAIT,
                'landscape': WD_ORIENT.LANDSCAPE,
            }
            #####
            print(obj_result)
            if len(obj_result) > 0:  # если есть данные то
                # code for docx

                document = Document()

                section = document.sections
                # 2481×3507

                for sec in section:
                    sec.orientation = orient_dict['portrait']
                    sec.page_width = Cm(29.8)
                    sec.page_height = Cm(21)

                font_styles = document.styles
                font_charstyle = font_styles.add_style(
                    'Head', WD_STYLE_TYPE.CHARACTER)
                font_object = font_charstyle.font
                font_object.size = Pt(14)
                font_object.name = 'Times New Roman'

                p = document.add_paragraph('')
                p.add_run('Отчёт по сотрудникам', style='Head').bold = True
                p.alignment = 1  # выравнивание: 0 - влево, 1 - центр

                # document.add_picture('monty-truth.png', width=Inches(1.25))
                thead_list = ['Дата', 'Фамилия, имя, отчество (при наличии) работника, прошедшего инструктаж по охране труда', 'профессия (должность) работника', 'число, месяц, год рождения работника', 'вид инструктажа по охране труда',
                              'Причина прохождения инструктажа по охране труда', 'ФИО, профессия работника, проводившего инструктаж', 'Наименование ЛПА, в объеме требований которого проведен инструктаж по охране труда']

                # заполнение шапки таблицы
                table = document.add_table(
                    rows=len(obj_result), cols=9, style='Table Grid')
                hdr_cells = table.rows[0].cells
                for i in range(len(thead_list)):
                    hdr_cells[i].text = thead_list[i]
                    hdr_cells[i].width = Cm(2.5)
                # ------

                # заполнение таблицы:
                # for i in range(len(obj_result)):
                #     for item in obj_result:
                #         row = table.add_row().cells
                #         row[i]

                # 89274420900 Дядя Володя Нурлат Черемуха

                dateNameFile = f'{timezone.localtime(timezone.now()).replace(hour=23, minute=59, second=0, microsecond=0)}.docx'
                response = HttpResponse(
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename=' + dateNameFile

                document.save(response)
                return response
            else:
                return HttpResponseBadRequest()

        if request.POST:  # для аякс условия

            if request.POST['id_date_start'] != '':  # если есть начальная дата
                try:
                    obj_result = res.objects.filter(
                        user__type_user=request.POST['id_type_user'],
                        instruction=request.POST['id_brief'],
                        user__groupStud_id=request.POST['id_group'],
                        quiz__id=request.POST['id_quiz'],
                        date_instruction__range=(request.POST['id_date_start'],
                                                 timezone.now()),
                        mark=request.POST['mark'],
                    )
                except res.DoesNotExist:
                    obj_result = []

            else:  # иначе без дат
                try:
                    obj_result = res.objects.filter(
                        user__type_user=request.POST['id_type_user'],
                        instruction=request.POST['id_brief'],
                        user__groupStud_id=request.POST['id_group'],
                        quiz__id=request.POST['id_quiz'],
                        mark=request.POST['mark'],
                    )
                except res.DoesNotExist:
                    obj_result = []

            if len(obj_result) > 0:

                data_list = []  # лист для данных после фильтра

                for row in obj_result:
                    data_list.append({
                        'surname': str(row.user.last_name),
                        'name': str(row.user.first_name),
                        'patro': str(row.user.patronymic),
                        'group': str(row.user.groupStud),  # группа
                        # тип пользователя
                        'type_user': str(row.user.type_user),
                        # тип пользователя (параметр в тестах)
                        'type_user_test': str(row.quiz.type_user),
                        # название инструктажа
                        'brief': str(row.instruction.name_instruction),
                        'quiz_name': str(row.quiz.name_test),
                        'date_target': str(row.quiz.date_target),
                        'date_passed': row.date_instruction,
                        'days_skiped': row.date_instruction_end,
                        'score': str(row.result),
                        'mark': str(row.mark),
                        'attempt': str(row.attempt),
                    })

                obj_res = data_list
                message = 'Успешно!'
                status = 'ok'

            else:
                message = 'Данные не найдены.'
                status = 'error'

            return JsonResponse({'result': obj_res, 'status': status, 'message': message}, safe=False)
    else:
        return render(request, "error.html")

    values = {
        'type_brief': obj_briefs,
        'type_users': obj_type_users,
        'groups': obj_groups,
        'quiz': obj_quizes,
        'users': obj_results_user,
    }

    return render(request, page_name, values)


#профиль
@login_required
class ProfileView(View):

    #профиль
    def showPage(request):
        page_name = 'profile.html'
        return render(request, page_name)

    #смена данных
    def changeData(request):
        if request.method == 'POST':
            data = {}  # for messages
            user = request.user

            if user.check_password(request.POST['password']):

                # проверка на наличие номера телефона в запросе
                if request.POST.get('phone_number'):  # если отправили номер телефона
                    if user.phone_number != request.POST['phone_number']:
                        form = ChangeNumberUser(request.POST)
                        if form.is_valid():
                            user.phone_number = request.POST['phone_number']
                            user.save()
                            data['message'] = f'{user.first_name}, Вы успешно сменили номер телефона!'
                            data['status'] = 'ok'
                            return JsonResponse(data)
                        else:
                            data['message'] = 'Введите правильный номер телефона!'
                            data['status'] = 'error'
                    else:
                        data['message'] = 'Вы ввели настоящий номер телефона!'
                        data['status'] = 'error'

                # проверка на почты в запросе
                if request.POST.get('email'):  # если отправили почту
                    if user.email != request.POST['email']:
                        form = ChangeEmailUser(request.POST)
                        if form.is_valid():
                            user.email = request.POST['email']
                            user.save()
                            data['message'] = f'{user.first_name}, Вы успешно сменили электронную почту!'
                            data['status'] = 'ok'
                            return JsonResponse(data)
                        else:
                            data['message'] = 'Введите корректную электронную почту!'
                            data['status'] = 'error'

                    else:
                        data['message'] = 'Новая почта совпадает с старым'
                        data['status'] = 'error'

                # смена пароля
                if request.POST.get('password2'):
                    if request.POST['password'] != request.POST['password2']:
                        form = ChangePasswordUser(request.POST)
                        if form.is_valid():
                            user.set_password(request.POST['password2'])
                            user.save()
                            data['message'] = f'{user.first_name}, Вы успешно сменили пароль!'
                            data['status'] = 'ok'
                            data['reload'] = 'go'
                            return JsonResponse(data)
                        else:
                            data['message'] = 'Новый пароль не соответствует по требованиям! '
                            data['status'] = 'error'

                    else:
                        data['message'] = 'Старый пароль соответствует новому'
                        data['status'] = 'error'

                # проверка на почты в запросе
                if request.FILES:  # если отправили почту
                    form = ChangeAvatarUser(request.FILES)
                    if form.is_valid():
                        user.avatar.delete()
                        user.avatar = request.FILES['avatar']
                        user.save()
                        data['message'] = f'{user.first_name}, Вы успешно сменили фотографию!'
                        data['status'] = 'ok'
                        return JsonResponse(data)
                    else:
                        data['message'] = 'Выберите корректную фотографию!'
                        data['status'] = 'error'

            else:
                data['message'] = 'Введите корректный пароль!'
                data['status'] = 'error'

            return JsonResponse(data)
    
    
    # пройденные инструктажи
    def passedBriefs(request):
        page_name = 'passed_inst.html'
        obj_res = res.objects.filter(
            user=request.user).order_by("-date_instruction")
        values = {'obj': obj_res}
        return render(request, page_name, values)


    #вкладка с деталями про пользователя
    def detail(request):
        page_name = 'profiile_detail_remaster.html'
        return render(request, page_name)


    #вкладка с возможностями
    def action(request):
        page_name = 'action.html'
        return render(request, page_name)


    #редактирование профиля
    def edit(request):
        page_name = 'profile_edit.html'
        return render(request, page_name)


    # сортировка результатов прохождения инструктажей
    def sorting(request):
        obj_res = None  # для результатов.
        if request.method == 'POST':  # Если метод пост
            if request.POST.get('filter'):  # Если существует элемент фильтра

                try:
                    obj_res = res.objects.filter(user=request.user).order_by(
                        request.POST['filter'])  # получение данных
                except res.DoesNotExist:
                    obj_res = []  # если не получилось, то

                if len(obj_res) > 0:  # длинна массива больше 0

                    data_list = []  # переменная лист для фильтра

                    for row in obj_res:
                        data_list.append({'name_brief': str(row.instruction.name_instruction), 'quiz_name': str(row.quiz.name_test), 'date_start': row.date_instruction, 'result': row.result, 'mark': row.mark})

                    obj_res = data_list

                    message = 'Успешно!'
                    status = 'ok'
                else:
                    message = 'Данные не найдены'
                    status = 'error'

                return JsonResponse({'result': obj_res, 'status': status, 'message': message})

        return HttpResponseBadRequest()



# инструктаж
@login_required
class BriefBrainView(View):


    # получение вопросов
    def questions(request, num, id):
        quiz = test.objects.get(
        instruction=id, type_user=request.user.type_user, pk=num)
        questions = []
        for q in quiz.get_questions():
            answers = []
            for a in q.get_answers():
                answers.append(a.text)
            questions.append({str(q): answers})
        return JsonResponse({
            'data': questions,
            'time': quiz.time
        })



    
    # сохранение теста
    def save(request, num, id):
        
        def getDaysPassed(first_date): # вычисление даты
            answer = timedelta.Timedelta(first_date - timezone.now())
            return abs(answer.total.days)
        


        def savetothedb(request): # сохранение в бд функция


            def saverequest(request, count_rows): # действие сохранения данных в бд
                count_rows.date_instruction = timezone.now()
                count_rows.date_instruction_end = request.brief_days_passed
                count_rows.result = request.brief_score
                count_rows.mark = request.brief_mark
                count_rows.attempt += 1 # прибавляем попытку
                count_rows.save()
                return True


            count_rows = res.objects.get(user = request.user, instruction_id = request.brief_instruction_id, quiz = request.brief_quiz) # получение строк


            if count_rows: # если есть уже строка
                return saverequest(request, count_rows)
            else: # создание
                res.objects.create( user=request.user,  instruction_id=request.brief_instruction_id,  quiz=request.brief_quiz,   date_instruction=timezone.now(), date_instruction_end = request.brief_days_passed, result= request.brief_score, mark=request.brief_mark,)
                return True
                
              


        if request.method == 'POST':
            user = request.user  # пользователь
            # получение нужного теста
            quiz = test.objects.get(pk=num, instruction=id,
                                    type_user=request.user.type_user)

            questions = []  # list с в вопросами
            data = request.POST
            data_ = dict(data.lists())  # в хороший список
            data_.pop('csrfmiddlewaretoken')

            for k in data_.keys():  # прогонка по вопросам
                quest = question.objects.get(name=k)
                questions.append(quest)  # Добавляем в список вопросов

            score = 0  # балл
            multiper = 100 / quiz.number_of_questions
            results = []
            correct_answer = None

            for q in questions:
                a_selected = request.POST.get(q.name)  # выбранные ответы

                if a_selected != "":  # если ответ не пустой, то
                    question_answers = answers.objects.filter(
                        question=q)  # получаем ответы по вопросу
                    for a in question_answers:
                        if a_selected == a.text:  # если выбранный ответ совпадает с ответов настоящим
                            if a.correct:
                                score += 1  # плюс отвеченный вопрос
                                correct_answer = a.text  # сохраняем правильный ответ
                        else:
                            if a.correct:
                                correct_answer = a.text
                    results.append(
                        {str(q): {'correct_answer': correct_answer, 'answered': a_selected}})
                else:
                    results.append({str(q): 'not answered'})

            countAnswers = score
            score_ = score * multiper

            # получение разницы в датах
            days_passed = getDaysPassed(quiz.date_target) # получение
            
            # заполнение запроса
            request.brief_instruction_id = quiz.instruction.id # номер инструктажа #
            request.brief_quiz = quiz # сам тест  #
            request.brief_days_passed = days_passed # разница между днями #
            request.brief_score = score_ # баллы # 
            request.brief_results = results # результаты
            request.brief_countAnswers = countAnswers # количество отвеченных вопросов
            # если человек набрал больше баллов, чем в условии теста, то сохраняем его и отправляем успешно
            #переписать 
            # найти строку по параметрам, если существует, то +1 попытка и сохранить, а если нет, то новую.
            if score_ >= quiz.required_score_to_pass:
                request.brief_mark = 'Сдан'
                savetothedb(request)
                return JsonResponse({'passed': True, 'score': score_, 'results': results, 'countAnswers': countAnswers})
            else:
                request.brief_mark = 'Не сдан'
                savetothedb(request)
                return JsonResponse({'passed': False, 'score': score_, 'results': results, 'countAnswers': countAnswers})


    # проверка насчет лекции
    def checkFile(request, id):
        data = {}
        if request.method == 'POST':
            obj_query = downloadInstructionsForTests.objects.filter(
                user=request.user, test=request.POST['quiz-pk']).count()
            if obj_query == 0:
                downloadInstructionsForTests.objects.create(
                    user=request.user,
                    test_id=request.POST['quiz-pk']
                )
                data['status'] = 'ok'
            else:
                data['status'] = 'warning'
            return JsonResponse(data)
        else:
            return HttpResponseBadRequest()
        

    # Проверка о прохождении теста
    def checkPassed(request, id): 
        data = {}
        if request.method == 'POST':
            obj_query_file = downloadInstructionsForTests.objects.filter(
                user=request.user, test=request.POST['quiz']).count()
            if obj_query_file != 0:
                passed_brief = res.objects.filter(
                    instruction=id, quiz=request.POST['quiz'], mark='Сдан', user=request.user).count()
                if passed_brief == 0:
                    data['status'] = True
                else:
                    data['message'] = 'Вы уже прошли данный инструктаж'
                    data['status'] = False
            else:
                data['message'] = 'Вы не изучили теорию!'
                data['status'] = False
        else:
            return HttpResponseBadRequest()
        return JsonResponse(data)



# получение страницы теста
@login_required  # обязательная авторизация
def testView(request, num, id):
    page_name = 'user_tests/user_tests_test.html'
    quiz = test.objects.get(
        instruction=id, type_user=request.user.type_user, pk=num)
    passed_brief = res.objects.filter(
        instruction=id, quiz=quiz, mark='Сдан', user=request.user).count()
    if passed_brief == 0:
        values = {
            'test': quiz,
        }
        return render(request, page_name, values)
    else:
        return HttpResponseBadRequest()
    


class BriefLayoutView(View):


    # def typeBrief(request):
    #     page_name = 'user_tests/user_tests_list.html'
    #     obj_themes = 
    #     return render(request, )

    @login_required
    def listBriefs(request):
        page_name = "user_tests/user_themes_tests_list.html"
        themesInstructions = inst.objects.all().order_by("-name_instruction")
        

        # сдедать пункт должники в журнале


        # выводить в журнале последний результат по тесту (дублирование человека не должно быть в журнале) лучше сделать через цикл проверку.
        # data = {}
        # параметры
        # user__type_user
        # instruction
        # user__groupStud_id
        # quiz__id
        # mark проверка без него
        # row - сама строка, если нет совпадений, то добавить саму строку
        # сделать проверку по этим параметрам в дата + сама строка
        # если есть совпадение, то сравнивать марк и менять его, если в дата строка равняется не сдал, а в новой строке сдал, то записать новую строку
        # в конце можно собрать все строки из дата ()

        # или добавить новое поле с количеством попыток


        for i in themesInstructions:
            i.counttest2 = test.get_need_instr(request, i.pk).count

        values = {
            'themes': themesInstructions,
        }

        return render(request, page_name, values)
    

    #страница с тестами
    @login_required 
    def tests(request, id):
        page_name = 'user_tests/user_tests_list.html'

        themesInstructions = inst.objects.get(id=id)  # model inst
        tests = test.get_need_instr(request, id)  # model test
        
        values = {
            'test': tests,
            'themes': themesInstructions,
        }

        return render(request, page_name, values)
